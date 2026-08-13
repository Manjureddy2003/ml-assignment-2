"""Build the editable DOCX and submission PDF from project artifacts.

Usage from the project root:
    python tools/build_submission_report.py

Expected output:
    Created report/submission_report.docx
    Created report/submission_report.pdf   (when LibreOffice is available)

To insert the mandatory lab evidence, save a genuine screenshot as:
    report/bits_virtual_lab_screenshot.png
and run this script again.
"""

from __future__ import annotations

import json
import math
from pathlib import Path
import shutil
import subprocess
import sys
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from sklearn.metrics import ConfusionMatrixDisplay, confusion_matrix

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
ARTIFACT_DIR = ROOT / "model" / "artifacts"
CONFIG_PATH = ROOT / "student_config.json"
METADATA_PATH = ARTIFACT_DIR / "metadata.json"
METRICS_PATH = ARTIFACT_DIR / "metrics.csv"
PREDICTIONS_PATH = ARTIFACT_DIR / "prediction_details.json"
TEST_DATA_PATH = ROOT / "test_data.csv"
SCREENSHOT_PATH = REPORT_DIR / "bits_virtual_lab_screenshot.png"
DOCX_PATH = REPORT_DIR / "submission_report.docx"
PDF_PATH = REPORT_DIR / "submission_report.pdf"
COMPARISON_CHART_PATH = REPORT_DIR / "model_comparison.png"
CONFUSION_CHART_PATH = REPORT_DIR / "winner_confusion_matrix.png"

NAVY = "17365D"
BLUE = "315BBF"
TEAL = "178A7A"
LIGHT_BLUE = "E9EEF8"
LIGHT_TEAL = "E8F5F2"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D9DEE7"
DARK_TEXT = "172033"
WHITE = "FFFFFF"
WARNING_FILL = "FFF4D6"
WARNING_BORDER = "D89A00"


def load_inputs() -> tuple[dict, dict, pd.DataFrame, dict, pd.DataFrame]:
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8"))
    config = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    metrics = pd.read_csv(METRICS_PATH)
    prediction_details = json.loads(PREDICTIONS_PATH.read_text(encoding="utf-8"))
    test_data = pd.read_csv(TEST_DATA_PATH)
    return metadata, config, metrics, prediction_details, test_data


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs: dict[str, str]) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn("w:" + key), str(edge_data[key]))


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn("w:" + margin_name))
        if node is None:
            node = OxmlElement("w:" + margin_name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def remove_paragraph_borders(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Liberation Sans"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("6B7280")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE, underline: bool = True) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    run_properties.append(color_element)
    if underline:
        underline_element = OxmlElement("w:u")
        underline_element.set(qn("w:val"), "single")
        run_properties.append(underline_element)
    new_run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def configure_document(document: Document, student_id: str) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(17)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(DARK_TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 28, NAVY),
        ("Subtitle", 13, TEAL),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 13.5, BLUE),
        ("Heading 3", 11.5, TEAL),
    ):
        style = styles[style_name]
        style.font.name = "Liberation Sans"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    for style_name in ("Title", "Subtitle"):
        style_element = styles[style_name].element
        p_pr = style_element.find(qn("w:pPr"))
        if p_pr is not None:
            p_bdr = p_pr.find(qn("w:pBdr"))
            if p_bdr is not None:
                p_pr.remove(p_bdr)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Machine Learning Assignment 2 | {student_id}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_para.runs[0].font.name = "Liberation Sans"
    footer_para.runs[0].font.size = Pt(8.5)
    footer_para.runs[0].font.color.rgb = RGBColor.from_string("6B7280")
    page_para = footer.add_paragraph()
    add_page_number(page_para)


def add_cover(document: Document, metadata: dict, config: dict) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Machine Learning\nAssignment 2")
    remove_paragraph_borders(title)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(metadata["project_title"])
    remove_paragraph_borders(subtitle)

    course = document.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run("M.Tech (AIML / DSE) | Work Integrated Learning Programmes")
    run.bold = True
    run.font.name = "Liberation Sans"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    document.add_paragraph()

    identity = document.add_table(rows=4, cols=2)
    identity.alignment = WD_TABLE_ALIGNMENT.CENTER
    identity.autofit = False
    identity.columns[0].width = Inches(1.7)
    identity.columns[1].width = Inches(4.8)
    identity_rows = [
        ("Student Name", config.get("student_name", "REPLACE WITH YOUR NAME")),
        ("Student ID", config.get("student_id", "REPLACE WITH YOUR BITS ID")),
        ("Dataset", metadata["dataset_name"]),
        ("Submission Date", "REPLACE WITH FINAL SUBMISSION DATE"),
    ]
    for row, (label, value) in zip(identity.rows, identity_rows):
        prevent_row_split(row)
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": MID_GRAY},
                bottom={"val": "single", "sz": "6", "color": MID_GRAY},
                left={"val": "single", "sz": "6", "color": MID_GRAY},
                right={"val": "single", "sz": "6", "color": MID_GRAY},
            )
        row.cells[0].paragraphs[0].runs[0].bold = True

    document.add_paragraph()
    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_TEAL)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "8", "color": TEAL},
        bottom={"val": "single", "sz": "8", "color": TEAL},
        left={"val": "single", "sz": "8", "color": TEAL},
        right={"val": "single", "sz": "8", "color": TEAL},
    )
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        "Six classification pipelines | 30 features | 569 instances | "
        "Interactive Streamlit deployment"
    )
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)

    document.add_page_break()


def add_section_heading(document: Document, number: str, title: str) -> None:
    heading = document.add_heading(f"{number}. {title}", level=1)
    heading.paragraph_format.page_break_before = False


def add_callout(document: Document, title: str, body: str, warning: bool = False) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    fill = WARNING_FILL if warning else LIGHT_BLUE
    border = WARNING_BORDER if warning else BLUE
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=150, end=150)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "8", "color": border},
        bottom={"val": "single", "sz": "8", "color": border},
        left={"val": "single", "sz": "8", "color": border},
        right={"val": "single", "sz": "8", "color": border},
    )
    para = cell.paragraphs[0]
    title_run = para.add_run(title + ": ")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(border)
    para.add_run(body)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_links_table(document: Document, config: dict) -> None:
    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)

    entries = [
        ("GitHub Repository", config.get("github_url", "REPLACE WITH YOUR GITHUB REPOSITORY URL")),
        ("Live Streamlit App", config.get("streamlit_url", "REPLACE WITH YOUR LIVE STREAMLIT URL")),
        ("Entry Point", "app.py"),
        ("Test Data", "test_data.csv"),
    ]
    for row, (label, value) in zip(table.rows, entries):
        prevent_row_split(row)
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        value_para = row.cells[1].paragraphs[0]
        if isinstance(value, str) and value.startswith("http"):
            add_hyperlink(value_para, value, value)
        else:
            value_para.add_run(str(value))
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=115, bottom=115, start=115, end=115)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "5", "color": MID_GRAY},
                bottom={"val": "single", "sz": "5", "color": MID_GRAY},
                left={"val": "single", "sz": "5", "color": MID_GRAY},
                right={"val": "single", "sz": "5", "color": MID_GRAY},
            )


def add_dataset_table(document: Document, metadata: dict) -> None:
    rows = [
        ("Dataset", metadata["dataset_name"]),
        ("Repository", "UCI Machine Learning Repository"),
        ("Instances", str(metadata["dataset_instances"])),
        ("Predictor Features", str(metadata["dataset_features"])),
        ("Task", "Binary classification"),
        ("Training Rows", str(metadata["train_rows"])),
        ("Test Rows", str(metadata["test_rows"])),
        ("Split", "Stratified 80:20 split"),
        ("Random State", str(metadata["random_state"])),
        ("Target Mapping", "0 = Benign, 1 = Malignant"),
        ("Positive Class", metadata["positive_class"]),
    ]
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    header = table.rows[0]
    header.cells[0].text = "Dataset Field"
    header.cells[1].text = "Value"
    set_repeat_table_header(header)
    for cell in header.cells:
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=100, bottom=100, start=110, end=110)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for label, value in rows:
        row = table.add_row()
        prevent_row_split(row)
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=95, bottom=95, start=110, end=110)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": MID_GRAY},
                bottom={"val": "single", "sz": "4", "color": MID_GRAY},
                left={"val": "single", "sz": "4", "color": MID_GRAY},
                right={"val": "single", "sz": "4", "color": MID_GRAY},
            )

    source_para = document.add_paragraph()
    source_para.add_run("Dataset source: ").bold = True
    add_hyperlink(source_para, metadata["dataset_source"], metadata["dataset_source"])


def add_model_details_table(document: Document) -> None:
    rows = [
        ("Logistic Regression", "Median imputation + standardization", "Balanced class weights; max_iter=5000"),
        ("Decision Tree", "Median imputation", "max_depth=5; min_samples_leaf=4; balanced weights"),
        ("K-Nearest Neighbors", "Median imputation + standardization", "k=7; distance weighting"),
        ("Gaussian Naive Bayes", "Median imputation + standardization", "var_smoothing=1e-9"),
        ("Random Forest", "Median imputation", "400 trees; min_samples_leaf=2; balanced weights"),
        ("Support Vector Machine", "Median imputation + standardization", "RBF kernel; C=2.0; probability=True; balanced weights"),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.65), Inches(2.25), Inches(2.6)]
    headers = ["Model", "Preprocessing", "Key Configuration"]
    for index, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        cell.text = text
        cell.width = widths[index]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=100, bottom=100, start=90, end=90)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for model, preprocessing, configuration in rows:
        row = table.add_row()
        prevent_row_split(row)
        values = [model, preprocessing, configuration]
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.text = value
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=95, bottom=95, start=90, end=90)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": MID_GRAY},
                bottom={"val": "single", "sz": "4", "color": MID_GRAY},
                left={"val": "single", "sz": "4", "color": MID_GRAY},
                right={"val": "single", "sz": "4", "color": MID_GRAY},
            )
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        row.cells[0].paragraphs[0].runs[0].bold = True


def add_metrics_table(document: Document, metrics: pd.DataFrame) -> None:
    headers = ["ML Model Name", "Accuracy", "AUC", "Precision", "Recall", "F1", "MCC"]
    widths = [Inches(1.65)] + [Inches(0.78)] * 6
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
        cell.text = header
        cell.width = widths[index]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=105, bottom=105, start=55, end=55)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8.4)
            run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])

    ranked = metrics.sort_values(by=["MCC", "F1", "AUC", "Accuracy"], ascending=False)
    winner_name = str(ranked.iloc[0]["Model"])
    for _, metric_row in metrics.iterrows():
        row = table.add_row()
        prevent_row_split(row)
        values = [str(metric_row["Model"])] + [
            f"{float(metric_row[column]):.4f}" for column in headers[1:]
        ]
        is_winner = str(metric_row["Model"]) == winner_name
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            cell.text = value
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            set_cell_margins(cell, top=90, bottom=90, start=45, end=45)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": MID_GRAY},
                bottom={"val": "single", "sz": "4", "color": MID_GRAY},
                left={"val": "single", "sz": "4", "color": MID_GRAY},
                right={"val": "single", "sz": "4", "color": MID_GRAY},
            )
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8.4)
                if is_winner:
                    run.bold = True
        if is_winner:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_TEAL)
        else:
            set_cell_shading(row.cells[0], LIGHT_GRAY)
            row.cells[0].paragraphs[0].runs[0].bold = True


def model_observations(metrics: pd.DataFrame) -> list[tuple[str, str]]:
    lookup = {row["Model"]: row for _, row in metrics.iterrows()}
    return [
        (
            "Logistic Regression",
            "Strong linear baseline after feature standardization. It achieved "
            f"{lookup['Logistic Regression']['Accuracy']:.4f} accuracy, "
            f"{lookup['Logistic Regression']['Recall']:.4f} recall, and "
            f"{lookup['Logistic Regression']['MCC']:.4f} MCC.",
        ),
        (
            "Decision Tree",
            "Lowest overall result on this split. The lower recall and MCC indicate "
            "that one depth-controlled tree generalized less effectively than the "
            "scaled and ensemble models.",
        ),
        (
            "K-Nearest Neighbors",
            "Good precision after scaling, with a strong MCC. It missed more malignant "
            "cases than Logistic Regression and SVM, which reduced recall and F1.",
        ),
        (
            "Gaussian Naive Bayes",
            "Very high AUC but lower thresholded recall and F1. Its probability ranking "
            "was strong, while the conditional-independence assumption limited the final "
            "class decisions.",
        ),
        (
            "Random Forest",
            "Best AUC at 0.9974 and perfect malignant-class precision on this test split. "
            "It produced no false-positive malignant predictions but missed three "
            "malignant rows.",
        ),
        (
            "Support Vector Machine",
            "Best overall balance under the MCC-first ranking. It achieved the highest "
            "accuracy, F1, and MCC, together with perfect precision and high recall.",
        ),
        (
            "Overall Winner",
            "Support Vector Machine. The ranking rule uses MCC first, then F1, AUC, "
            "and Accuracy. Random Forest remains the AUC winner.",
        ),
    ]


def add_observations_table(document: Document, metrics: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(4.75)
    headers = ["ML Model Name", "Observation about model performance"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=105, bottom=105, start=100, end=100)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])

    for model_name, observation in model_observations(metrics):
        row = table.add_row()
        prevent_row_split(row)
        row.cells[0].text = model_name
        row.cells[1].text = observation
        set_cell_shading(row.cells[0], LIGHT_GRAY if model_name != "Overall Winner" else LIGHT_TEAL)
        if model_name == "Overall Winner":
            set_cell_shading(row.cells[1], LIGHT_TEAL)
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=105, bottom=105, start=105, end=105)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": MID_GRAY},
                bottom={"val": "single", "sz": "4", "color": MID_GRAY},
                left={"val": "single", "sz": "4", "color": MID_GRAY},
                right={"val": "single", "sz": "4", "color": MID_GRAY},
            )


def create_comparison_chart(metrics: pd.DataFrame) -> None:
    chart = metrics.set_index("Model")[["Accuracy", "F1", "MCC"]]
    axis = chart.plot(kind="bar", figsize=(10.5, 5.2))
    axis.set_ylim(0.7, 1.01)
    axis.set_ylabel("Score")
    axis.set_xlabel("")
    axis.set_title("Model comparison on the bundled test set")
    axis.grid(axis="y", alpha=0.25)
    plt.xticks(rotation=25, ha="right")
    plt.tight_layout()
    plt.savefig(COMPARISON_CHART_PATH, dpi=180, bbox_inches="tight")
    plt.close()


def create_confusion_chart(
    metadata: dict, prediction_details: dict, test_data: pd.DataFrame
) -> None:
    winner = metadata["overall_winner"]
    predictions = np.asarray(prediction_details[winner]["predictions"], dtype=int)
    target = test_data["target"].to_numpy(dtype=int)
    matrix = confusion_matrix(target, predictions, labels=[0, 1])
    figure, axis = plt.subplots(figsize=(5.5, 4.5))
    display = ConfusionMatrixDisplay(
        confusion_matrix=matrix, display_labels=["Benign", "Malignant"]
    )
    display.plot(ax=axis, values_format="d", colorbar=False)
    axis.set_title(f"{winner}: confusion matrix")
    figure.tight_layout()
    figure.savefig(CONFUSION_CHART_PATH, dpi=180, bbox_inches="tight")
    plt.close(figure)


def add_screenshot_section(document: Document) -> None:
    document.add_page_break()
    document.add_heading("8. BITS Virtual Lab Execution Evidence", level=1)
    document.add_paragraph(
        "The assignment requires one screenshot showing successful execution on BITS "
        "Virtual Lab. The screenshot must be captured by the student in the required "
        "environment."
    )
    if SCREENSHOT_PATH.exists():
        document.add_picture(str(SCREENSHOT_PATH), width=Inches(6.5))
        caption = document.add_paragraph("Figure: Assignment execution on BITS Virtual Lab")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].italic = True
        caption.runs[0].font.size = Pt(9)
    else:
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_shading(cell, WARNING_FILL)
        set_cell_margins(cell, top=520, bottom=520, start=220, end=220)
        set_cell_border(
            cell,
            top={"val": "dashed", "sz": "12", "color": WARNING_BORDER},
            bottom={"val": "dashed", "sz": "12", "color": WARNING_BORDER},
            left={"val": "dashed", "sz": "12", "color": WARNING_BORDER},
            right={"val": "dashed", "sz": "12", "color": WARNING_BORDER},
        )
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first = para.add_run("INSERT GENUINE BITS VIRTUAL LAB SCREENSHOT HERE")
        first.bold = True
        first.font.size = Pt(15)
        first.font.color.rgb = RGBColor.from_string(WARNING_BORDER)
        para.add_run("\n\nSave it as report/bits_virtual_lab_screenshot.png and rebuild the report.")
    document.add_page_break()


def add_repository_tree(document: Document) -> None:
    tree = """project-folder/
|-- app.py
|-- requirements.txt
|-- README.md
|-- test_data.csv
|-- student_config.json
|-- data/
|-- model/
|   |-- train_models.py
|   |-- model_registry.py
|   `-- artifacts/ (six saved classifiers and metric files)
|-- utils/
|-- tests/
|-- tools/build_submission_report.py
`-- report/submission_report.docx and submission_report.pdf"""
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=130, bottom=130, start=150, end=150)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "5", "color": MID_GRAY},
        bottom={"val": "single", "sz": "5", "color": MID_GRAY},
        left={"val": "single", "sz": "5", "color": MID_GRAY},
        right={"val": "single", "sz": "5", "color": MID_GRAY},
    )
    para = cell.paragraphs[0]
    run = para.add_run(tree)
    run.font.name = "Liberation Mono"
    run.font.size = Pt(8.6)


def add_bullet_list(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def build_document(
    metadata: dict,
    config: dict,
    metrics: pd.DataFrame,
    prediction_details: dict,
    test_data: pd.DataFrame,
) -> Document:
    create_comparison_chart(metrics)
    create_confusion_chart(metadata, prediction_details, test_data)

    document = Document()
    configure_document(document, config.get("student_id", "REPLACE WITH YOUR BITS ID"))
    add_cover(document, metadata, config)

    add_section_heading(document, "1", "Mandatory Submission Links")
    document.add_paragraph(
        "The final submission PDF must contain clickable links to the complete GitHub "
        "repository and the deployed Streamlit application."
    )
    add_links_table(document, config)
    if not str(config.get("github_url", "")).startswith("http") or not str(
        config.get("streamlit_url", "")
    ).startswith("http"):
        add_callout(
            document,
            "Action required",
            "Replace both link placeholders after the repository and app are live. "
            "Rebuild the report before submission.",
            warning=True,
        )

    add_section_heading(document, "2", "Problem Statement")
    document.add_paragraph(
        "The objective is to implement multiple classification models on one public "
        "dataset, evaluate every model with Accuracy, AUC, Precision, Recall, F1 Score, "
        "and Matthews Correlation Coefficient, and demonstrate the saved models through "
        "an interactive Streamlit web application."
    )
    add_callout(
        document,
        "Model-count interpretation",
        "The brief names five mandatory models but also refers to six models. This "
        "project includes every named model and adds a Support Vector Machine as the "
        "sixth classifier.",
    )

    add_section_heading(document, "3", "Dataset Description")
    document.add_paragraph(
        "The Breast Cancer Wisconsin (Diagnostic) dataset is used because it satisfies "
        "the minimum size constraints and supports a clear binary classification task. "
        "The predictors are numeric measurements derived from digitized images of "
        "breast-mass cell nuclei."
    )
    add_dataset_table(document, metadata)

    document.add_heading("Feature organization", level=2)
    document.add_paragraph(
        "The 30 predictors contain ten base measurements represented as mean values, "
        "standard-error values, and worst/extreme values. The full feature dictionary "
        "is included in data/feature_dictionary.csv."
    )

    add_section_heading(document, "4", "Modeling Methodology")
    add_bullet_list(
        document,
        [
            "Normalize all predictor names to lowercase snake_case.",
            "Remap the target to 0 = Benign and 1 = Malignant.",
            "Create one stratified 80:20 train-test split with random_state=42.",
            "Use median imputation in every saved pipeline so uploaded CSV files can contain missing numeric cells.",
            "Standardize features for distance-based, linear, probabilistic, and kernel models.",
            "Train all six classifiers on the same training rows and evaluate them on the same 114 test rows.",
            "Save each fitted pipeline as a joblib artifact for Streamlit deployment.",
        ],
    )
    document.add_heading("Implemented models", level=2)
    add_model_details_table(document)

    add_section_heading(document, "5", "Evaluation Results")
    document.add_paragraph(
        "All metrics use Malignant (1) as the positive class. AUC is calculated from the "
        "model probability or decision score rather than the final thresholded class."
    )
    add_metrics_table(document, metrics)
    document.add_paragraph()
    document.add_picture(str(COMPARISON_CHART_PATH), width=Inches(6.55))
    caption = document.add_paragraph("Figure 1. Accuracy, F1, and MCC comparison")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)

    document.add_heading("Winner confusion matrix", level=2)
    document.add_picture(str(CONFUSION_CHART_PATH), width=Inches(4.8))
    caption = document.add_paragraph(
        f"Figure 2. {metadata['overall_winner']} predictions on the bundled test set"
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)

    add_section_heading(document, "6", "Model Performance Observations")
    add_observations_table(document, metrics)

    add_section_heading(document, "7", "Streamlit Application and Repository")
    document.add_heading("Required application features", level=2)
    add_bullet_list(
        document,
        [
            "CSV upload control with bundled test_data.csv as the default evaluation file.",
            "Model-selection dropdown covering all six saved classifiers.",
            "Metric cards for Accuracy, AUC, Precision, Recall, F1, and MCC.",
            "Confusion matrix and classification report for labeled test data.",
            "All-model comparison table and performance chart.",
            "Row-level class predictions and malignant-class probabilities.",
            "Prediction CSV download control.",
            "Feature validation, target conversion, and missing-value handling.",
        ],
    )
    document.add_heading("Repository structure", level=2)
    add_repository_tree(document)

    document.add_heading("Local execution", level=2)
    code_lines = [
        "python -m pip install -r requirements.txt",
        "python -m model.train_models",
        "streamlit run app.py",
    ]
    code_table = document.add_table(rows=1, cols=1)
    code_cell = code_table.cell(0, 0)
    set_cell_shading(code_cell, LIGHT_GRAY)
    set_cell_margins(code_cell, top=120, bottom=120, start=150, end=150)
    code_run = code_cell.paragraphs[0].add_run("\n".join(code_lines))
    code_run.font.name = "Liberation Mono"
    code_run.font.size = Pt(9)
    document.add_paragraph(
        "Expected result: a browser opens with the upload control, model dropdown, six "
        "metric cards, confusion matrix, classification report, comparison table, and "
        "prediction download button."
    )

    add_screenshot_section(document)

    add_section_heading(document, "9", "Conclusion")
    document.add_paragraph(
        "The project completes the full machine learning workflow: dataset preparation, "
        "consistent model training, multi-metric evaluation, artifact persistence, "
        "interactive visualization, and deployment packaging. On the bundled test split, "
        f"{metadata['overall_winner']} is the overall winner under the stated MCC-first "
        "ranking, while Random Forest achieves the highest AUC."
    )
    add_callout(
        document,
        "Important limitation",
        "This is an educational classification demonstration and must not be used as a "
        "medical diagnostic system.",
        warning=True,
    )

    add_section_heading(document, "10", "Final Submission Checklist")
    checklist_items = [
        "GitHub repository link opens and contains complete source code.",
        "requirements.txt, README.md, and test_data.csv are present in the repository root.",
        "All six model artifacts are committed under model/artifacts/.",
        "Live Streamlit app opens in a private/incognito browser window.",
        "CSV upload, model selection, metrics, confusion matrix, and classification report work.",
        "README content is represented in this PDF.",
        "One genuine BITS Virtual Lab execution screenshot is included.",
        "Student name, BITS ID, GitHub URL, and Streamlit URL are finalized.",
    ]
    for item in checklist_items:
        para = document.add_paragraph()
        run = para.add_run("[ ] ")
        run.bold = True
        para.add_run(item)

    document.add_heading("Appendix A. Reproducibility Notes", level=1)
    document.add_paragraph(
        "Run python -m model.train_models to regenerate the train/test CSV files, all "
        "six model artifacts, metrics.csv, and metadata.json. Run pytest -q after "
        "installing requirements-dev.txt; the expected result is 6 passed."
    )
    document.add_paragraph(
        "The included metrics are tied to the supplied dataset, the stratified split, "
        "random_state=42, and the recorded model settings. Changing any of these items "
        "can change the comparison table and the overall winner."
    )

    return document


def convert_to_pdf(docx_path: Path, output_dir: Path) -> Path | None:
    executable = shutil.which("libreoffice") or shutil.which("soffice")
    if executable is None:
        return None
    command = [
        executable,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(docx_path),
    ]
    result = subprocess.run(command, capture_output=True, text=True, check=False)
    generated = output_dir / (docx_path.stem + ".pdf")
    if result.returncode != 0 or not generated.exists():
        print(result.stdout)
        print(result.stderr, file=sys.stderr)
        return None
    return generated


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    metadata, config, metrics, prediction_details, test_data = load_inputs()
    document = build_document(metadata, config, metrics, prediction_details, test_data)
    document.save(DOCX_PATH)
    print(f"Created {DOCX_PATH.relative_to(ROOT)}")
    pdf = convert_to_pdf(DOCX_PATH, REPORT_DIR)
    if pdf is not None:
        print(f"Created {pdf.relative_to(ROOT)}")
    else:
        print("LibreOffice PDF conversion was unavailable; export the DOCX manually.")


if __name__ == "__main__":
    main()
